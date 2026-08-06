"""
Turns an uploaded supplier catalog (PDF / DOCX / image) into content Claude
can read: either extracted text, or base64 images for pages that are mostly
pictures/tables (common in product brochures).

PDFs use iter_pdf_pages(): a generator that yields ONE page at a time
(render -> caller uses it -> caller discards it -> next page). This keeps
at most one rendered page image in memory at a time, regardless of how many
pages the catalog has -- important on small-memory hosting tiers.
"""
import base64
import io
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from PIL import Image

from app import config

PDF_RENDER_DPI = 72  # low DPI keeps each page image small; catalogs stay readable


def extract_content(filepath: str) -> dict:
    """
    Returns {"text": str, "images": [base64_png, ...]}
    Used for non-PDF files (DOCX/image/Excel/CSV -- all small enough to
    process in one shot). For PDFs, use iter_pdf_pages() instead.
    """
    ext = Path(filepath).suffix.lower()

    if ext in (".docx", ".doc"):
        return _extract_docx(filepath)
    elif ext in (".png", ".jpg", ".jpeg", ".webp"):
        return _extract_image(filepath)
    elif ext in (".xlsx", ".csv"):
        return _extract_tabular(filepath)
    else:
        raise ValueError(f"Unsupported file type for extract_content: {ext} (use iter_pdf_pages for PDFs)")


def iter_pdf_pages(filepath: str):
    """
    Generator: yields (page_number, text, image_b64_or_None) one page at a
    time. Only one rendered page image exists in memory at any point.

    image_b64 is None for pages with enough extractable text that a
    rendered image isn't needed.

    If you break out of iterating early, call the generator's .close()
    method (or use it in a try/finally) so the underlying PDF file handle
    gets released via the `finally` block below.
    """
    doc = fitz.open(filepath)
    try:
        total_pages = min(len(doc), config.MAX_PDF_PAGES)
        for i in range(total_pages):
            page = doc[i]
            text = page.get_text()
            image_b64 = None

            if len(text.strip()) < 40:
                pix = page.get_pixmap(dpi=PDF_RENDER_DPI)
                img_bytes = pix.tobytes("png")
                image_b64 = base64.b64encode(img_bytes).decode()
                pix = None
                img_bytes = None

            yield i, text, image_b64
    finally:
        doc.close()


def _extract_docx(filepath: str) -> dict:
    doc = Document(filepath)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return {"text": "\n".join(parts), "images": []}


def _extract_image(filepath: str) -> dict:
    img = Image.open(filepath).convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return {"text": "", "images": [base64.b64encode(buf.getvalue()).decode()]}


def _extract_tabular(filepath: str) -> dict:
    import csv
    ext = Path(filepath).suffix.lower()
    rows = []
    if ext == ".csv":
        with open(filepath, newline="", encoding="utf-8", errors="ignore") as f:
            rows = list(csv.reader(f))
    else:
        from openpyxl import load_workbook
        wb = load_workbook(filepath, read_only=True, data_only=True)
        ws = wb.active
        for row in ws.iter_rows(values_only=True):
            rows.append(["" if c is None else str(c) for c in row])
        wb.close()
    text = "\n".join(",".join(r) for r in rows)
    return {"text": text, "images": []}
